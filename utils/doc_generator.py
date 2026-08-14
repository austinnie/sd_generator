# tools/utils/doc_generator.py
"""Word 文档生成工具"""

import os
import sys
import glob
from datetime import datetime

CURRENT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if CURRENT_DIR not in sys.path:
    sys.path.insert(0, CURRENT_DIR)


def generate_word_doc(output_dir: str, folder_name: str, reviews: list) -> bool:
    """
    生成 Word 文档
    返回是否成功
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("   ⚠️ 未安装 python-docx，跳过 Word 生成")
        return False
    
    try:
        # 获取当前子文件夹的所有图片
        valid_images = []
        for ext in ['.jpg', '.jpeg', '.png']:
            valid_images.extend(glob.glob(os.path.join(output_dir, f"*{ext}")))
        valid_images = [img for img in valid_images if os.path.exists(img)]
        
        if not valid_images:
            print(f"      ⚠️ 未找到有效图片，跳过 Word 生成")
            return False
        
        doc = Document()
        
        # 设置标题
        title = doc.add_heading(f"【{folder_name} 作品合辑】", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加元数据
        meta = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_after = Inches(0.2)
        
        # 遍历图片和点评
        for idx, img_path in enumerate(valid_images):
            # 获取对应的点评
            review_text = f"【作品 {idx+1}】\n（请在此处微调你的专属评论）"
            if idx < len(reviews):
                clean_review = reviews[idx].replace("【", "").replace("】", "").strip()
                review_text = f"【作品 {idx+1}】\n{clean_review}"
            
            # 插入图片
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(5.5))
            except Exception as e:
                print(f"      ⚠️ Word 插入图片失败: {e}")
            
            # 插入鉴赏文字
            review_p = doc.add_paragraph(review_text)
            review_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            review_p.paragraph_format.space_before = Inches(0.1)
            review_p.paragraph_format.space_after = Inches(0.3)
        
        # 保存
        docx_file = os.path.join(output_dir, "公众号草稿.docx")
        doc.save(docx_file)
        print(f"      📄 已生成 Word 文档：{os.path.basename(docx_file)}")
        return True
        
    except Exception as e:
        print(f"      ⚠️ Word 文档生成失败：{e}")
        return False


def generate_text_summary(output_dir: str, folder_name: str, reviews: list) -> bool:
    """
    生成文本摘要
    返回是否成功
    """
    try:
        summary_file = os.path.join(output_dir, "点评.txt")
        with open(summary_file, "w", encoding="utf-8") as f:
            f.write(f"【{folder_name} AI 作品鉴赏合辑】\n")
            f.write(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write(f"本集共收录 {len(reviews)} 件 AI 视觉创作：\n\n")
            
            for review in reviews:
                f.write(f"{review}\n\n")
        
        print(f"      📝 已生成备份 Txt 文档：{os.path.basename(summary_file)}")
        return True
    except Exception as e:
        print(f"      ⚠️ Txt 备份文档写入失败：{e}")
        return False